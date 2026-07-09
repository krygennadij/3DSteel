# -*- coding: utf-8 -*-
"""Генерация отчётов: Word (.docx) и Excel (.xlsx) с нативными графиками."""
import io
import math
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import openpyxl
from openpyxl.chart import Reference, ScatterChart, Series
from openpyxl.chart.axis import ChartLines
from openpyxl.chart.layout import Layout, ManualLayout
from openpyxl.chart.marker import Marker
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.styles import (Alignment, Border, Font, PatternFill, Side,
                              numbers)
from openpyxl.utils import get_column_letter
import pandas as pd

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from calc import FireCalcResult, MPA_TO_KGF_CM2


# ═══════════════════════════════════════════════════════════════════════════
# Вспомогательные функции — изображения (matplotlib)
# ═══════════════════════════════════════════════════════════════════════════

def _capacity_chart_png(res: FireCalcResult, dpi: int = 150) -> bytes:
    """График несущей способности vs момент от нагрузки (PNG)."""
    t_arr = res.load_capacity["Время, мин"].to_numpy(dtype=float)
    cap   = res.load_capacity["Несущая способность, кНм"].to_numpy(dtype=float)
    mom   = res.applied_moment_value
    limit = res.fire_limit_minute

    fig, ax = plt.subplots(figsize=(10, 5))

    finite_cap = cap[np.isfinite(cap)]
    y_max = float(finite_cap[0]) * 1.08 if len(finite_cap) > 0 else 100.0

    ax.plot(t_arr, cap, color="#e60000", linewidth=2.0,
            label="Несущая способность, кНм")
    ax.axhline(mom, color="#0055cc", linewidth=2.0, linestyle="--",
               label=f"Момент от нагрузки = {mom:.3f} кНм")

    if limit is not None and limit > 0 and limit < len(t_arr):
        ax.axvline(limit, color="black", linewidth=1.0, linestyle=":")
        x_right = float(t_arr[-1])
        # Подпись у линии предела: слева/справа от неё, чтобы не залезать
        # в легенду (закреплённую в верхнем правом углу)
        if limit > 0.6 * x_right:
            label_x, ha = limit - 0.3, "right"
        else:
            label_x, ha = limit + 0.3, "left"
        ax.text(label_x, y_max * 0.55,
                f"tпред = {limit} мин",
                fontsize=10, va="top", ha=ha,
                bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="none"))

    ax.set_xlim(left=0, right=float(t_arr[-1]))
    ax.set_ylim(bottom=0, top=y_max)
    ax.set_xlabel("Время, мин", fontsize=12)
    ax.set_ylabel("Момент, кНм", fontsize=12)
    ax.grid(True, alpha=0.25, linewidth=0.8)
    ax.legend(loc="upper right", fontsize=10)
    ax.tick_params(labelsize=10)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _section_png(b: float, h: float, t: float, s: float,
                 heated_sides: set, dpi: int = 120) -> bytes:
    """Схема поперечного сечения двутавра с выделением обогреваемых сторон."""
    hh, hb, hs = h / 2, b / 2, s / 2

    xs = [-hb,  hb,  hb,  hs,  hs,  hb,  hb, -hb, -hb, -hs, -hs, -hb, -hb]
    ys = [ hh,  hh, hh-t, hh-t, -(hh-t), -(hh-t), -hh, -hh, -(hh-t), -(hh-t), hh-t, hh-t, hh]

    fig, ax = plt.subplots(figsize=(4, 5))
    ax.fill(xs, ys, color="#d8d8d8", linewidth=1.2, edgecolor="#444")

    ORANGE, GRAY, LW = "#ff8800", "#bbbbbb", 4.5
    side_paths = {
        "bottom": ([-hb,  hb], [-hh,  -hh]),
        "top":    ([-hb,  hb], [ hh,   hh]),
        "right":  ([ hb,  hb,  hs,  hs,  hb,  hb],
                   [ hh, hh-t, hh-t, -(hh-t), -(hh-t), -hh]),
        "left":   ([-hb, -hb, -hs, -hs, -hb, -hb],
                   [-hh, -(hh-t), -(hh-t), hh-t, hh-t, hh]),
    }
    for side, (px, py) in side_paths.items():
        color = ORANGE if side in heated_sides else GRAY
        ax.plot(px, py, color=color, linewidth=LW, solid_capstyle="round")

    ax.set_aspect("equal")
    ax.axis("off")
    pad = max(hb, hh) * 0.25
    ax.set_xlim(-hb - pad, hb + pad)
    ax.set_ylim(-hh - pad, hh + pad)
    ax.set_title(f"b={b:.0f}  h={h:.0f}  t={t:.1f}  s={s:.1f} мм",
                 fontsize=9, pad=6)

    legend_els = [
        mpatches.Patch(color=ORANGE, label="Обогревается"),
        mpatches.Patch(color=GRAY,   label="Не обогревается"),
    ]
    ax.legend(handles=legend_els, loc="lower center",
              bbox_to_anchor=(0.5, -0.08), ncol=2, fontsize=8, frameon=False)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════
# Вспомогательные функции — Word
# ═══════════════════════════════════════════════════════════════════════════

_HDR_FILL = "2C5F8A"
_ALT_FILL = "EAF2FB"


def _cell_fill(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _cell_text(cell, text: str, bold=False, size_pt=12,
               color_hex: str = None, align="center") -> None:
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(_ru(str(text)))
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color_hex:
        r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_df_table(doc: Document, df: pd.DataFrame,
                  col_widths_cm: list = None) -> None:
    """Добавляет DataFrame как таблицу в документ Word."""
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Table Grid"

    # Ширина столбцов
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Заголовок
    for j, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        _cell_fill(cell, _HDR_FILL)
        _cell_text(cell, col_name, bold=True, size_pt=12, color_hex="FFFFFF")

    # Данные
    for i, (_, row) in enumerate(df.iterrows()):
        fill = _ALT_FILL if i % 2 == 1 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if fill != "FFFFFF":
                _cell_fill(cell, fill)
            col_name = df.columns[j]
            if isinstance(val, (float, np.floating)):
                if np.isnan(val) or np.isinf(val):
                    text = "—"
                elif col_name == "Время, мин":
                    text = f"{val:.0f}"
                else:
                    text = f"{val:.3f}"
            else:
                text = str(val)
            _cell_text(cell, text, size_pt=12)


def _result_box(doc: Document, label: str, value: str,
                fill_hex: str = "EAF2FB") -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _cell_fill(tbl.rows[0].cells[0], fill_hex)
    _cell_text(tbl.rows[0].cells[0], label, bold=True, size_pt=12)
    _cell_text(tbl.rows[0].cells[1], value, size_pt=12, align="center")


# ── Формулы с подстрочными индексами ────────────────────────────────────────

_DECIMAL_RE = re.compile(r"(?<=\d)\.(?=\d)")


def _ru(text: str) -> str:
    """Заменяет десятичную точку на запятую (принятый в РФ разделитель)."""
    return _DECIMAL_RE.sub(",", text)


def _n(text: str) -> list:
    return [(_ru(text), False)]


def _s(text: str) -> list:
    return [(text, True)]


def _sym(base: str, sub: str = "") -> list:
    return [(base, False), (sub, True)] if sub else [(base, False)]


def _cat(*parts) -> list:
    out = []
    for part in parts:
        out.extend(part)
    return out


class _Frac:
    """Токен-дробь: числитель/знаменатель — списки токенов (текст, признак_подстрочного)."""
    __slots__ = ("num_tokens", "den_tokens")

    def __init__(self, num_tokens: list, den_tokens: list):
        self.num_tokens = num_tokens
        self.den_tokens = den_tokens


def _frac(num_tokens: list, den_tokens: list) -> list:
    """Дробь num/den, отображаемая в Word как настоящая дробь (числитель над чертой)."""
    return [_Frac(num_tokens, den_tokens)]


# ── Настоящие дроби (OMML) ──────────────────────────────────────────────────

def _omath_run(text: str, size_pt: float):
    """Один текстовый run внутри математической зоны Word (m:r)."""
    r = OxmlElement("m:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Cambria Math")
    rfonts.set(qn("w:hAnsi"), "Cambria Math")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(size_pt * 2))))
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def _omath_tokens(tokens: list, size_pt: float) -> list:
    """Строит элементы OMML (m:r, либо m:sSub для пар текст+подстрочный) из токенов."""
    elems = []
    for text, is_sub in tokens:
        if is_sub and elems is not None and len(elems) > 0 and elems[-1].tag == qn("m:r"):
            base_r = elems.pop()
            ssub = OxmlElement("m:sSub")
            e = OxmlElement("m:e")
            e.append(base_r)
            sub = OxmlElement("m:sub")
            sub.append(_omath_run(text, size_pt))
            ssub.append(e)
            ssub.append(sub)
            elems.append(ssub)
        else:
            elems.append(_omath_run(text, size_pt))
    return elems


def _omath_frac(frac: _Frac, size_pt: float):
    """m:oMath с одной дробью (числитель над знаменателем, с чертой)."""
    om = OxmlElement("m:oMath")
    f = OxmlElement("m:f")
    num = OxmlElement("m:num")
    for e in _omath_tokens(frac.num_tokens, size_pt):
        num.append(e)
    den = OxmlElement("m:den")
    for e in _omath_tokens(frac.den_tokens, size_pt):
        den.append(e)
    f.append(num)
    f.append(den)
    om.append(f)
    return om


def _feq(doc: Document, tokens: list, size: float = 16.5, align: str = "center",
         italic: bool = False) -> None:
    """Строка-формула из токенов: (текст, признак_подстрочного) либо _Frac (дробь)."""
    p = doc.add_paragraph()
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    for tok in tokens:
        if isinstance(tok, _Frac):
            p._p.append(_omath_frac(tok, size))
            continue
        text, is_sub = tok
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.italic = italic
        if is_sub:
            run.font.subscript = True


def _mixed_paragraph(doc: Document, tokens: list, size: float = 15) -> None:
    """Обычный (не центрированный) абзац с подстрочными фрагментами и дробями."""
    p = doc.add_paragraph()
    for tok in tokens:
        if isinstance(tok, _Frac):
            p._p.append(_omath_frac(tok, size))
            continue
        text, is_sub = tok
        run = p.add_run(text)
        run.font.size = Pt(size)
        if is_sub:
            run.font.subscript = True


def _fnum(v, nd: int = 2) -> str:
    if v is None or (isinstance(v, (float, np.floating)) and not np.isfinite(v)):
        return "—"
    return _ru(f"{v:.{nd}f}")


def _add_methodology_section(doc: Document, section_no: int, res: FireCalcResult,
                             b_mm: float, h_mm: float, t_mm: float, s_mm: float,
                             m_kgm: float, load_kg: float, length_m: float) -> None:
    """Раздел с формулами методики и численной подстановкой для одной
    показательной минуты пожара."""

    n_last = len(res.load_capacity) - 1
    limit = res.fire_limit_minute
    if limit is not None and limit > 0:
        demo = limit
        demo_note = f"момент наступления предела огнестойкости"
    elif limit == 0:
        demo = 0
        demo_note = "начальный момент времени (предел огнестойкости не достигается позже)"
    else:
        demo = min(15, n_last)
        demo_note = "иллюстративный момент времени"

    st_row   = res.strength_table.iloc[demo]
    zone_row = res.compressed_zone.iloc[demo]
    rn_row   = res.normative_resistance.iloc[demo]
    arm_row  = res.lever_arms.iloc[demo]
    mo_row   = res.bending_moments.iloc[demo]
    cap_t    = res.load_capacity["Несущая способность, кНм"].iloc[demo]
    mom_val  = res.applied_moment_value

    yld_lower_0 = res.strength_table["Предел текучести нижней полки"].iloc[0]
    yld_web_0   = res.strength_table["Предел текучести стенки"].iloc[0]
    yld_upper_0 = res.strength_table["Предел текучести верхней полки"].iloc[0]

    temp_lower = st_row["Температура нижней полки, ℃"]
    yld_lower  = st_row["Предел текучести нижней полки"]
    k_lower    = st_row["Коэффициент снижения предела текучести нижней полки"]
    temp_web   = st_row["Температура стенки, ℃"]
    yld_web    = st_row["Предел текучести стенки"]
    k_web      = st_row["Коэффициент снижения предела текучести стенки"]
    temp_upper = st_row["Температура верхней полки, ℃"]
    yld_upper  = st_row["Предел текучести верхней полки"]
    k_upper    = st_row["Коэффициент снижения предела текучести верхней полки"]

    a_flange = zone_row["Показатель сжатой зоны при x < a"]
    a_web    = zone_row["Показатель сжатой зоны при x > a"]
    in_web   = bool(np.isfinite(a_flange) and a_flange > t_mm)
    a_used   = a_web if in_web else a_flange

    rn_lower = rn_row["Нормативное сопротивление нижней полки, кгс/см²"]
    rn_web   = rn_row["Нормативное сопротивление стенки, кгс/см²"]
    rn_upper = rn_row["Нормативное сопротивление верхней полки, кгс/см²"]

    arm_tensile_lower     = arm_row["Плечо равнодействующей силы растяжения в нижней полке, мм"]
    arm_tensile_web       = arm_row["Плечо равнодействующей силы растяжения в нижней части стенки, мм"]
    arm_compression_web   = arm_row["Плечо равнодействующей силы сжатия в верхней части стенки, мм"]
    arm_compression_upper = arm_row["Плечо равнодействующей силы сжатия в верхней полке, мм"]

    moment_lower     = mo_row["Изгибающий момент в нижней полке, кНм"]
    moment_web       = mo_row["Изгибающий момент в нижней части стенки, кНм"]
    moment_upper_web = mo_row["Изгибающий момент в верхней части стенки, кНм"]
    moment_upper     = mo_row["Изгибающий момент в верхней полке, кНм"]

    tensile_lower     = rn_lower * b_mm * t_mm * 0.01
    if in_web:
        tensile_web       = rn_web * s_mm * (h_mm - a_used - t_mm) * 0.01
        compression_web   = rn_web * s_mm * (a_used - t_mm) * 0.01
        compression_upper = rn_upper * t_mm * b_mm * 0.01
    else:
        tensile_web       = rn_web * s_mm * (h_mm - 2 * t_mm) * 0.01
        compression_web   = rn_upper * b_mm * (t_mm - a_used) * 0.01
        compression_upper = rn_upper * a_used * b_mm * 0.01

    doc.add_heading(f"{section_no}. Методика расчёта: пример с численной подстановкой", level=1)
    doc.add_paragraph(_ru(
        f"Ниже показан порядок расчёта, заложенный в программу, с подстановкой "
        f"реальных чисел для сечения b×h×t×s = {b_mm:.0f}×{h_mm:.0f}×{t_mm:.1f}×{s_mm:.1f} мм "
        f"— {demo_note} (t = {demo} мин). Расчёт для остальных минут выполняется по тем же "
        f"формулам; полные результаты — в таблицах ниже."
    )).runs[0].font.size = Pt(15)

    # 1. Снижение прочности стали
    doc.add_heading(f"{section_no}.1. Снижение прочности стали при нагреве", level=2)
    doc.add_paragraph(
        "Коэффициент снижения предела текучести каждого участка сечения — отношение "
        "предела текучести стали при текущей температуре к пределу текучести при 20 °C:"
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(
        _sym("γ", "x"), _n(" = "),
        _frac(_cat(_sym("R", "y,x"), _n("(t)")),
              _cat(_sym("R", "y,x"), _n("(20 °C)"))),
    ), italic=True)
    doc.add_paragraph(
        f"При t = {demo} мин температуры участков: нижняя полка — {temp_lower:.0f} °C, "
        f"стенка — {temp_web:.0f} °C, верхняя полка — {temp_upper:.0f} °C."
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(_sym("γ", "н"),  _n(" = "),
                   _frac(_n(f"{yld_lower:.1f}"), _n(f"{yld_lower_0:.1f}")),
                   _n(f" = {k_lower:.3f}")))
    _feq(doc, _cat(_sym("γ", "ст"), _n(" = "),
                   _frac(_n(f"{yld_web:.1f}"), _n(f"{yld_web_0:.1f}")),
                   _n(f" = {k_web:.3f}")))
    _feq(doc, _cat(_sym("γ", "в"),  _n(" = "),
                   _frac(_n(f"{yld_upper:.1f}"), _n(f"{yld_upper_0:.1f}")),
                   _n(f" = {k_upper:.3f}")))

    # 2. Нормативное сопротивление
    doc.add_heading(f"{section_no}.2. Нормативное сопротивление стали", level=2)
    _feq(doc, _cat(_sym("R", "n,x"), _n(" = "), _sym("γ", "x"), _n(" · R"),
                   _sym("y", "x"), _n("(20 °C) · 10.197")), italic=True)
    _feq(doc, _cat(_sym("R", "n,н"),
                   _n(f" = {k_lower:.3f} · {yld_lower_0:.1f} · 10.197 = {rn_lower:.1f} кгс/см²")))
    _feq(doc, _cat(_sym("R", "n,ст"),
                   _n(f" = {k_web:.3f} · {yld_web_0:.1f} · 10.197 = {rn_web:.1f} кгс/см²")))
    _feq(doc, _cat(_sym("R", "n,в"),
                   _n(f" = {k_upper:.3f} · {yld_upper_0:.1f} · 10.197 = {rn_upper:.1f} кгс/см²")))

    # 3. Положение нейтральной оси
    doc.add_heading(f"{section_no}.3. Положение нейтральной оси", level=2)
    doc.add_paragraph(
        "Сначала проверяется, укладывается ли граница сжатой зоны в толщину полки (x ≤ t):"
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(
        _n("a = "),
        _frac(
            _cat(_sym("γ", "в"), _n("·b·t + "), _sym("γ", "ст"),
                 _n("·s·h − 2"), _sym("γ", "ст"), _n("·s·t + "), _sym("γ", "н"), _n("·b·t")),
            _cat(_n("2"), _sym("γ", "в"), _n("·b")),
        ),
    ), italic=True)
    _feq(doc, _cat(
        _n("a = "),
        _frac(
            _n(f"{k_upper:.3f}·{b_mm:.0f}·{t_mm:.1f} + {k_web:.3f}·{s_mm:.1f}·{h_mm:.0f} − "
               f"2·{k_web:.3f}·{s_mm:.1f}·{t_mm:.1f} + {k_lower:.3f}·{b_mm:.0f}·{t_mm:.1f}"),
            _n(f"2·{k_upper:.3f}·{b_mm:.0f}"),
        ),
        _n(f" = {_fnum(a_flange)} мм"),
    ))
    if in_web:
        doc.add_paragraph(_ru(
            f"Так как a = {_fnum(a_flange)} мм > t = {t_mm:.1f} мм, нейтральная ось лежит "
            f"в стенке — показатель сжатой зоны пересчитывается по формуле для стенки:"
        )).runs[0].font.size = Pt(15)
        _feq(doc, _cat(
            _n("a = "),
            _frac(
                _cat(_sym("γ", "ст"), _n("·h·s − "), _sym("γ", "в"),
                     _n("·t·b + "), _sym("γ", "н"), _n("·t·b")),
                _cat(_n("2"), _sym("γ", "ст"), _n("·s")),
            ),
        ), italic=True)
        _feq(doc, _cat(
            _n("a = "),
            _frac(
                _n(f"{k_web:.3f}·{h_mm:.0f}·{s_mm:.1f} − {k_upper:.3f}·{t_mm:.1f}·{b_mm:.0f} + "
                   f"{k_lower:.3f}·{t_mm:.1f}·{b_mm:.0f}"),
                _n(f"2·{k_web:.3f}·{s_mm:.1f}"),
            ),
            _n(f" = {_fnum(a_web)} мм"),
        ))
    else:
        doc.add_paragraph(_ru(
            f"Так как a = {_fnum(a_flange)} мм ≤ t = {t_mm:.1f} мм, нейтральная ось лежит "
            f"в пределах полки — показатель сжатой зоны a = {_fnum(a_flange)} мм принимается "
            f"без пересчёта."
        )).runs[0].font.size = Pt(15)

    # 4. Усилия
    doc.add_heading(f"{section_no}.4. Усилия растяжения и сжатия", level=2)
    doc.add_paragraph(
        "Каждое усилие — произведение нормативного сопротивления участка на его "
        "площадь в сжатой/растянутой зоне (площадь в мм² переводится в см² "
        "множителем 0,01):"
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(_sym("N", "р.н"), _n(" = "), _sym("R", "n,н"), _n(" · b · t · 0.01")), italic=True)
    _feq(doc, _cat(_sym("N", "р.н"), _n(
        f" = {rn_lower:.1f} · {b_mm:.0f} · {t_mm:.1f} · 0.01 = {_fnum(tensile_lower, 1)} кгс"
    )))
    if in_web:
        _feq(doc, _cat(_sym("N", "р.ст"), _n(" = "), _sym("R", "n,ст"),
                       _n(" · s · (h − a − t) · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "р.ст"), _n(
            f" = {rn_web:.1f} · {s_mm:.1f} · ({h_mm:.0f} − {_fnum(a_used)} − {t_mm:.1f}) · 0.01 "
            f"= {_fnum(tensile_web, 1)} кгс"
        )))
        _feq(doc, _cat(_sym("N", "сж.ст"), _n(" = "), _sym("R", "n,ст"),
                       _n(" · s · (a − t) · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "сж.ст"), _n(
            f" = {rn_web:.1f} · {s_mm:.1f} · ({_fnum(a_used)} − {t_mm:.1f}) · 0.01 "
            f"= {_fnum(compression_web, 1)} кгс"
        )))
        _feq(doc, _cat(_sym("N", "сж.в"), _n(" = "), _sym("R", "n,в"),
                       _n(" · t · b · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "сж.в"), _n(
            f" = {rn_upper:.1f} · {t_mm:.1f} · {b_mm:.0f} · 0.01 = {_fnum(compression_upper, 1)} кгс"
        )))
    else:
        _feq(doc, _cat(_sym("N", "р.ст"), _n(" = "), _sym("R", "n,ст"),
                       _n(" · s · (h − 2t) · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "р.ст"), _n(
            f" = {rn_web:.1f} · {s_mm:.1f} · ({h_mm:.0f} − 2·{t_mm:.1f}) · 0.01 "
            f"= {_fnum(tensile_web, 1)} кгс"
        )))
        _feq(doc, _cat(_sym("N", "сж.ст"), _n(" = "), _sym("R", "n,в"),
                       _n(" · b · (t − a) · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "сж.ст"), _n(
            f" = {rn_upper:.1f} · {b_mm:.0f} · ({t_mm:.1f} − {_fnum(a_used)}) · 0.01 "
            f"= {_fnum(compression_web, 1)} кгс"
        )))
        _feq(doc, _cat(_sym("N", "сж.в"), _n(" = "), _sym("R", "n,в"),
                       _n(" · a · b · 0.01")), italic=True)
        _feq(doc, _cat(_sym("N", "сж.в"), _n(
            f" = {rn_upper:.1f} · {_fnum(a_used)} · {b_mm:.0f} · 0.01 = {_fnum(compression_upper, 1)} кгс"
        )))

    # 5. Плечи
    doc.add_heading(f"{section_no}.5. Плечи равнодействующих сил", level=2)
    doc.add_paragraph(
        "Плечи измеряются от найденной нейтральной оси до центра тяжести каждой зоны:"
    ).runs[0].font.size = Pt(15)
    if in_web:
        _feq(doc, _cat(
            _n("h − a − "), _frac(_n("t"), _n("2")),
            _n(f" = {h_mm:.0f} − {_fnum(a_used)} − "), _frac(_n(f"{t_mm:.1f}"), _n("2")),
            _n(f" = {_fnum(arm_tensile_lower)} мм  (плечо "), _sym("N", "р.н"), _n(")"),
        ))
        _feq(doc, _cat(
            _frac(_n("h − a − t"), _n("2")), _n(" = "),
            _frac(_n(f"{h_mm:.0f} − {_fnum(a_used)} − {t_mm:.1f}"), _n("2")),
            _n(f" = {_fnum(arm_tensile_web)} мм  (плечо "), _sym("N", "р.ст"), _n(")"),
        ))
        _feq(doc, _cat(
            _frac(_n("a − t"), _n("2")), _n(" = "),
            _frac(_n(f"{_fnum(a_used)} − {t_mm:.1f}"), _n("2")),
            _n(f" = {_fnum(arm_compression_web)} мм  (плечо "), _sym("N", "сж.ст"), _n(")"),
        ))
        _feq(doc, _cat(
            _n("a − "), _frac(_n("t"), _n("2")),
            _n(f" = {_fnum(a_used)} − "), _frac(_n(f"{t_mm:.1f}"), _n("2")),
            _n(f" = {_fnum(arm_compression_upper)} мм  (плечо "), _sym("N", "сж.в"), _n(")"),
        ))
    else:
        _feq(doc, _cat(
            _n("h − "), _frac(_n("t"), _n("2")), _n(" − a"),
            _n(f" = {h_mm:.0f} − "), _frac(_n(f"{t_mm:.1f}"), _n("2")),
            _n(f" − {_fnum(a_used)} = {_fnum(arm_tensile_lower)} мм  (плечо "),
            _sym("N", "р.н"), _n(")"),
        ))
        _feq(doc, _cat(
            _frac(_n("h"), _n("2")), _n(" − a"),
            _n(f" = {_fnum(h_mm/2, 1)} − {_fnum(a_used)}"),
            _n(f" = {_fnum(arm_tensile_web)} мм  (плечо "), _sym("N", "р.ст"), _n(")"),
        ))
        _feq(doc, _cat(
            _frac(_n("t − a"), _n("2")), _n(" = "),
            _frac(_n(f"{t_mm:.1f} − {_fnum(a_used)}"), _n("2")),
            _n(f" = {_fnum(arm_compression_web)} мм  (плечо "), _sym("N", "сж.ст"), _n(")"),
        ))
        _feq(doc, _cat(
            _frac(_n("a"), _n("2")), _n(" = "),
            _frac(_n(f"{_fnum(a_used)}"), _n("2")),
            _n(f" = {_fnum(arm_compression_upper)} мм  (плечо "), _sym("N", "сж.в"), _n(")"),
        ))

    # 6. Изгибающие моменты и несущая способность
    doc.add_heading(f"{section_no}.6. Изгибающие моменты и несущая способность", level=2)
    doc.add_paragraph(
        "Момент от каждого усилия — произведение усилия (кгс) на его плечо (мм); "
        "множитель 0,00001 переводит кгс·мм в кН·м (g ≈ 10 м/с²):"
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(_sym("M", "р.н"), _n(
        f" = {_fnum(tensile_lower, 1)} · {_fnum(arm_tensile_lower)} · 0.00001 "
        f"= {_fnum(moment_lower, 3)} кНм"
    )))
    _feq(doc, _cat(_sym("M", "р.ст"), _n(
        f" = {_fnum(tensile_web, 1)} · {_fnum(arm_tensile_web)} · 0.00001 "
        f"= {_fnum(moment_web, 3)} кНм"
    )))
    _feq(doc, _cat(_sym("M", "сж.ст"), _n(
        f" = {_fnum(compression_web, 1)} · {_fnum(arm_compression_web)} · 0.00001 "
        f"= {_fnum(moment_upper_web, 3)} кНм"
    )))
    _feq(doc, _cat(_sym("M", "сж.в"), _n(
        f" = {_fnum(compression_upper, 1)} · {_fnum(arm_compression_upper)} · 0.00001 "
        f"= {_fnum(moment_upper, 3)} кНм"
    )))
    _feq(doc, _cat(_sym("M", "нес"), _n(" = "),
                   _sym("M", "р.н"), _n(" + "), _sym("M", "р.ст"), _n(" + "),
                   _sym("M", "сж.ст"), _n(" + "), _sym("M", "сж.в"), _n(" = "),
                   _n(f"{_fnum(moment_lower,3)} + {_fnum(moment_web,3)} + "
                      f"{_fnum(moment_upper_web,3)} + {_fnum(moment_upper,3)} = "
                      f"{_fnum(cap_t, 2)} кНм")))

    # 7. Момент от нагрузки
    doc.add_heading(f"{section_no}.7. Момент от нагрузки", level=2)
    doc.add_paragraph(
        "Момент в середине пролёта от сосредоточенной нагрузки P и от собственного "
        "веса балки q = m·g (не зависит от времени пожара):"
    ).runs[0].font.size = Pt(15)
    _feq(doc, _cat(
        _sym("M", "нагр"), _n(" = "),
        _frac(_n("P·L"), _n("4")), _n("·9.80665·10⁻³ + "),
        _frac(_n("q·L²"), _n("8")), _n("·0.01"),
    ), italic=True)
    if m_kgm is not None:
        term1 = (load_kg * length_m / 4) * 9.80665e-3
        term2 = (m_kgm * length_m ** 2 / 8) * 0.01
        _feq(doc, _cat(
            _sym("M", "нагр"), _n(" = "),
            _frac(_n(f"{load_kg:.1f}·{length_m:.2f}"), _n("4")), _n("·9.80665·10⁻³ + "),
            _frac(_n(f"{m_kgm:.1f}·{length_m:.2f}²"), _n("8")), _n("·0.01"),
            _n(f" = {term1:.3f} + {term2:.3f} = {mom_val:.3f} кНм"),
        ))
    else:
        _feq(doc, _cat(_sym("M", "нагр"), _n(f" = {mom_val:.3f} кНм")))

    # 8. Предел огнестойкости
    doc.add_heading(f"{section_no}.8. Предел огнестойкости", level=2)
    _mixed_paragraph(doc, _cat(
        _n("Предел огнестойкости — последняя целая минута, для которой "),
        _sym("M", "нес"), _n("(t) ещё превышает "), _sym("M", "нагр"),
        _n(f" = {mom_val:.3f} кНм. На {demo}-й минуте "), _sym("M", "нес"),
        _n(f" = {_fnum(cap_t, 2)} кНм. Полная динамика по всем минутам — в таблице "
           f"«Несущая способность и момент от нагрузки» ниже."),
    ))


# ═══════════════════════════════════════════════════════════════════════════
# Генерация Word-отчёта
# ═══════════════════════════════════════════════════════════════════════════

def make_word_report(
    res: FireCalcResult,
    doc_name: str,
    profile_key: str,
    grade: str,
    load_kg: float,
    length_m: float,
    temp_source: str = "Встроенные данные (без ОГЗ)",
    b_mm: float = None,
    h_mm: float = None,
    t_mm: float = None,
    s_mm: float = None,
    m_kgm: float = None,
    heated_sides: set = None,
    geom_df: pd.DataFrame = None,
) -> bytes:
    """Возвращает байты .docx с подробным отчётом."""
    doc = Document()

    # ── Глобальный стиль ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(16.5)
    doc.styles["Heading 1"].font.size = Pt(21)
    doc.styles["Heading 2"].font.size = Pt(19.5)

    # ── Поля страницы ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(1.5)

    # ── Титульный блок ───────────────────────────────────────────────────────
    h = doc.add_heading("", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("РАСЧЁТ ПРЕДЕЛА ОГНЕСТОЙКОСТИ\nСТАЛЬНОЙ ИЗГИБАЕМОЙ БАЛКИ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(27)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "с учётом неравномерного прогрева двутаврового сечения\n"
        "по методике СП 2.13130.2020"
    )
    r.font.size = Pt(18)
    r.italic = True

    doc.add_paragraph()

    # ── 1. Исходные данные ───────────────────────────────────────────────────
    doc.add_heading("1. Исходные данные", level=1)

    rows_data = [
        ("Нормативный документ",     doc_name),
        ("Профиль двутавра",          profile_key),
        ("Марка стали",               grade),
        ("Сосредоточенная нагрузка P", f"{load_kg:.1f} кг"),
        ("Длина пролёта L",            f"{length_m:.2f} м"),
        ("Источник температур",        temp_source),
    ]
    if b_mm is not None:
        rows_data += [
            ("Высота сечения h",     f"{h_mm:.0f} мм"),
            ("Ширина полки b",        f"{b_mm:.0f} мм"),
            ("Толщина полки t",       f"{t_mm:.1f} мм"),
            ("Толщина стенки s",      f"{s_mm:.1f} мм"),
        ]
    if m_kgm is not None:
        rows_data.append(("Погонная масса",  f"{m_kgm:.1f} кг/м"))

    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(rows_data):
        fill = _ALT_FILL if i % 2 == 1 else "FFFFFF"
        _cell_fill(tbl.rows[i].cells[0], fill)
        _cell_fill(tbl.rows[i].cells[1], fill)
        _cell_text(tbl.rows[i].cells[0], lbl, bold=True, size_pt=12)
        _cell_text(tbl.rows[i].cells[1], val, size_pt=12)

    doc.add_paragraph()

    # ── 2. Результаты расчёта ────────────────────────────────────────────────
    doc.add_heading("2. Результаты расчёта", level=1)

    limit = res.fire_limit_minute
    cap0  = res.load_capacity["Несущая способность, кНм"].iloc[0]
    mom   = res.applied_moment_value

    if limit is None:
        limit_str = f"> {int(res.load_capacity['Время, мин'].iloc[-1])} мин"
        conclusion = "Несущая способность превышает момент от нагрузки на всём расчётном периоде."
    elif limit == 0:
        limit_str = "< 1 мин"
        conclusion = "Несущей способности не хватает уже без нагрева. Увеличьте сечение или уменьшите нагрузку."
    else:
        limit_str = f"{limit} мин"
        conclusion = (
            f"На {limit}-й минуте пожара несущая способность снижается до уровня "
            f"момента от нагрузки ({mom:.3f} кНм)."
        )

    res_data = [
        ("Предел огнестойкости",                limit_str),
        ("Несущая способность при t = 0, кНм",  f"{cap0:.2f}"),
        ("Момент от нагрузки M, кНм",            f"{mom:.3f}"),
    ]
    r_tbl = doc.add_table(rows=len(res_data), cols=2)
    r_tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(res_data):
        _cell_fill(r_tbl.rows[i].cells[0], "D6E4F0")
        _cell_text(r_tbl.rows[i].cells[0], lbl, bold=True, size_pt=12)
        _cell_text(r_tbl.rows[i].cells[1], val, size_pt=12, align="center")

    doc.add_paragraph()
    p_conc = doc.add_paragraph()
    r_conc = p_conc.add_run(_ru(conclusion))
    r_conc.italic = True
    r_conc.font.size = Pt(15)

    doc.add_paragraph()

    # ── 3. График несущей способности ────────────────────────────────────────
    doc.add_heading("3. График несущей способности", level=1)

    chart_bytes = _capacity_chart_png(res)
    doc.add_picture(io.BytesIO(chart_bytes), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap_note = doc.add_paragraph(
        "Рисунок 1 — Несущая способность балки при пожаре (сплошная красная линия) "
        "и момент от нагрузки (синяя пунктирная линия). "
        "Вертикальная линия — предел огнестойкости."
    )
    cap_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_note.runs[0].font.size = Pt(13.5)
    cap_note.runs[0].italic = True

    doc.add_paragraph()

    # ── 4. Схема сечения ─────────────────────────────────────────────────────
    _section_num = 4
    if b_mm is not None:
        doc.add_heading(f"{_section_num}. Схема поперечного сечения", level=1)
        hs_set = heated_sides if heated_sides is not None else {"bottom", "left", "right"}
        sec_bytes = _section_png(b_mm, h_mm, t_mm, s_mm, hs_set)
        doc.add_picture(io.BytesIO(sec_bytes), width=Inches(3.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        sides_map = {"bottom": "нижняя", "top": "верхняя",
                     "left": "левая", "right": "правая"}
        sides_str = ", ".join(sides_map.get(s, s) for s in sorted(hs_set))
        sec_note = doc.add_paragraph(
            f"Рисунок 2 — Поперечное сечение двутавра {profile_key}. "
            f"Обогреваемые стороны: {sides_str}."
        )
        sec_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sec_note.runs[0].font.size = Pt(13.5)
        sec_note.runs[0].italic = True
        doc.add_paragraph()
        _section_num += 1

    # ── Методика расчёта с численной подстановкой ────────────────────────────
    if b_mm is not None:
        _add_methodology_section(doc, _section_num, res,
                                 b_mm, h_mm, t_mm, s_mm, m_kgm,
                                 load_kg, length_m)
        doc.add_paragraph()
        _section_num += 1

    # ── Параметры огнезащиты ──────────────────────────────────────────────
    if geom_df is not None:
        doc.add_heading(f"{_section_num}. Параметры огнезащиты", level=1)
        _add_df_table(doc, geom_df)
        doc.add_paragraph()
        _section_num += 1

    # ── Подробные таблицы расчёта ─────────────────────────────────────────
    doc.add_heading(f"{_section_num}. Подробные таблицы расчёта", level=1)
    doc.add_paragraph(
        "Расчёт ведётся пошагово для каждой минуты: температуры трёх участков сечения → "
        "коэффициенты снижения прочности → нормативное сопротивление → "
        "положение нейтральной оси → усилия → изгибающие моменты → несущая способность."
    ).runs[0].font.size = Pt(15)
    doc.add_paragraph()

    steps = [
        ("Температуры и коэффициенты снижения предела текучести",
         res.strength_table),
        ("Положение нейтральной оси (показатель сжатой зоны)",
         res.compressed_zone),
        ("Нормативное сопротивление стали, кгс/см²",
         res.normative_resistance),
        ("Усилия растяжения и сжатия",
         res.efforts),
        ("Плечи равнодействующих сил, мм",
         res.lever_arms),
        ("Изгибающие моменты от усилий, кНм",
         res.bending_moments),
        ("Несущая способность и момент от нагрузки",
         pd.concat([res.load_capacity.set_index("Время, мин"),
                    res.applied_moment.set_index("Время, мин")],
                   axis=1).reset_index()),
    ]
    for i, (title, df) in enumerate(steps, 1):
        doc.add_heading(f"{_section_num}.{i}. {title}", level=2)
        _add_df_table(doc, df.round(3))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════
# Генерация Excel-отчёта
# ═══════════════════════════════════════════════════════════════════════════

_HDR_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
_HDR_FILL_XL = PatternFill("solid", fgColor="2C5F8A")
_ALT_FILL_XL = PatternFill("solid", fgColor="EAF2FB")
_BODY_FONT = Font(name="Calibri", size=10)
_CENTER    = Alignment(horizontal="center", vertical="center", wrap_text=True)
_LEFT      = Alignment(horizontal="left",   vertical="center", wrap_text=True)
_BORDER    = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"),  bottom=Side(style="thin"),
)


def _xl_write_df(ws, df: pd.DataFrame, start_row: int = 1,
                 start_col: int = 1) -> int:
    """Записывает DataFrame в лист openpyxl начиная с (start_row, start_col).
    Возвращает номер строки после последней строки данных."""
    # Заголовок
    for j, col_name in enumerate(df.columns, start_col):
        cell = ws.cell(row=start_row, column=j, value=col_name)
        cell.font   = _HDR_FONT
        cell.fill   = _HDR_FILL_XL
        cell.alignment = _CENTER
        cell.border = _BORDER

    # Данные
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row, start_col):
            r = start_row + 1 + i
            if isinstance(val, (float, np.floating)):
                if np.isnan(val) or np.isinf(val):
                    v = None
                else:
                    v = round(float(val), 4)
            elif isinstance(val, (int, np.integer)):
                v = int(val)
            else:
                v = val
            cell = ws.cell(row=r, column=j, value=v)
            cell.font      = _BODY_FONT
            cell.alignment = _CENTER
            cell.border    = _BORDER
            if i % 2 == 1:
                cell.fill = _ALT_FILL_XL

    # Авто-ширина столбцов
    for j in range(start_col, start_col + len(df.columns)):
        col_letter = get_column_letter(j)
        max_len = max(
            len(str(df.columns[j - start_col])),
            *(len(str(df.iloc[i, j - start_col])) for i in range(min(len(df), 50))),
        )
        ws.column_dimensions[col_letter].width = min(max_len + 2, 30)

    return start_row + 1 + len(df)


def _xl_section_title(ws, text: str, row: int, ncols: int) -> None:
    cell = ws.cell(row=row, column=1, value=text)
    cell.font = Font(name="Calibri", bold=True, size=11, color="1F3864")
    cell.alignment = _CENTER
    ws.merge_cells(start_row=row, start_column=1,
                   end_row=row, end_column=ncols)


def _nice_major_unit(max_value: float, target_ticks: int = 15):
    """Возвращает (шаг делений, верхняя граница оси), при которых и 0, и
    max_value гарантированно попадают на подписанное деление: шаг подбирается
    близким к target_ticks делений, а граница оси округляется вверх до
    ближайшего кратного шагу (если max_value само не делится нацело)."""
    if max_value <= 0:
        return 1, 1
    raw = max_value / target_ticks
    candidates = (1, 2, 5, 10, 15, 20, 25, 30, 50, 100, 150, 200, 250, 500, 1000)
    divisors = [c for c in candidates if c <= max_value and max_value % c == 0
                and 0.4 * raw <= c <= 2.5 * raw]
    if divisors:
        step = min(divisors, key=lambda c: abs(c - raw))
    else:
        step = next((c for c in candidates if c >= raw), candidates[-1])
    axis_max = math.ceil(max_value / step) * step
    return step, axis_max


def _make_capacity_chart(ws_data, n_rows: int) -> ScatterChart:
    """Создаёт нативный Excel ScatterChart (Точечная с гладкими кривыми)."""
    chart = ScatterChart()
    chart.title  = None
    chart.style  = 10
    chart.scatterStyle = "smooth"
    chart.x_axis.title = "Время, мин"
    chart.x_axis.axPos = "b"
    chart.x_axis.scaling.min = 0
    chart.y_axis.title = "Момент, кНм"
    chart.y_axis.axPos = "l"
    chart.y_axis.scaling.min = 0
    chart.height = 14
    chart.width  = 22

    # Ось X — время (столбец A), используется как реальные числовые значения
    x_ref = Reference(ws_data, min_col=1, min_row=2, max_row=n_rows + 1)

    # Серия 1: несущая способность (столбец B)
    cap_ref = Reference(ws_data, min_col=2, max_col=2,
                        min_row=1, max_row=n_rows + 1)
    s0 = Series(cap_ref, x_ref, title_from_data=True)
    s0.marker = Marker(symbol="none")
    s0.graphicalProperties.line.solidFill = "E60000"
    s0.graphicalProperties.line.width = 25000   # 1/12700 pt → ~2pt
    s0.smooth = True
    chart.series.append(s0)

    # Серия 2: момент от нагрузки (столбец C)
    mom_ref = Reference(ws_data, min_col=3, max_col=3,
                        min_row=1, max_row=n_rows + 1)
    s1 = Series(mom_ref, x_ref, title_from_data=True)
    s1.marker = Marker(symbol="none")
    s1.graphicalProperties.line.solidFill = "0055CC"
    s1.graphicalProperties.line.width = 25000
    s1.graphicalProperties.line.dashDot = "dash"
    s1.smooth = True
    chart.series.append(s1)

    # Основные линии сетки по обеим осям (в т.ч. вертикальные)
    chart.x_axis.majorGridlines = ChartLines()
    chart.y_axis.majorGridlines = ChartLines()
    chart.x_axis.delete = False
    chart.y_axis.delete = False

    return chart


def _make_comparison_chart_xl(ws_data, n_rows: int, series_colors: list) -> ScatterChart:
    """Нативный Excel ScatterChart для сравнения нескольких сценариев ОГЗ.

    Столбец A — время; столбцы B..(B+len(series_colors)-1) — несущая способность
    каждого сценария (цвета из series_colors); последний столбец — момент от
    нагрузки (общий, пунктирная синяя линия)."""
    chart = ScatterChart()
    chart.title  = None
    chart.style  = 10
    chart.scatterStyle = "smooth"
    chart.x_axis.title = "Время, мин"
    chart.x_axis.axPos = "b"
    _dtick, _axis_max = _nice_major_unit(n_rows - 1)
    chart.x_axis.scaling.min = 0
    chart.x_axis.scaling.max = _axis_max
    chart.x_axis.majorUnit = _dtick
    chart.y_axis.title = "Момент, кНм"
    chart.y_axis.axPos = "l"
    chart.y_axis.scaling.min = 0
    chart.height = 14
    chart.width  = 22

    # Область построения — фиксированные поля, чтобы заголовок оси Y и
    # легенда не наезжали на подписи делений (Excel сам не всегда даёт
    # графику достаточно места).
    chart.layout = Layout(
        manualLayout=ManualLayout(
            layoutTarget="inner", xMode="edge", yMode="edge",
            x=0.09, y=0.03, w=0.88, h=0.82,
        )
    )

    x_ref = Reference(ws_data, min_col=1, min_row=2, max_row=n_rows + 1)

    for i, color in enumerate(series_colors):
        col = 2 + i
        ref = Reference(ws_data, min_col=col, max_col=col, min_row=1, max_row=n_rows + 1)
        s = Series(ref, x_ref, title_from_data=True)
        s.marker = Marker(symbol="none")
        s.graphicalProperties.line.solidFill = color
        s.graphicalProperties.line.width = 25000
        s.smooth = True
        chart.series.append(s)

    mom_col = 2 + len(series_colors)
    mom_ref = Reference(ws_data, min_col=mom_col, max_col=mom_col,
                        min_row=1, max_row=n_rows + 1)
    s_mom = Series(mom_ref, x_ref, title_from_data=True)
    s_mom.marker = Marker(symbol="none")
    s_mom.graphicalProperties.line.solidFill = "0055CC"
    s_mom.graphicalProperties.line.width = 25000
    s_mom.graphicalProperties.line.dashDot = "dash"
    s_mom.smooth = True
    chart.series.append(s_mom)

    chart.x_axis.majorGridlines = ChartLines()
    chart.y_axis.majorGridlines = ChartLines()
    chart.x_axis.delete = False
    chart.y_axis.delete = False

    # Легенда — компактным блоком в правом верхнем углу поверх графика,
    # с белой подложкой и рамкой, чтобы не сжимать область построения.
    chart.legend.position = "r"
    chart.legend.overlay = True
    chart.legend.layout = Layout(
        manualLayout=ManualLayout(
            xMode="edge", yMode="edge", x=0.615, y=0.03, w=0.35, h=0.20,
        )
    )
    chart.legend.spPr = GraphicalProperties(solidFill="FFFFFF")
    chart.legend.spPr.line.solidFill = "000000"

    return chart


def make_excel_report(
    res: FireCalcResult,
    doc_name: str,
    profile_key: str,
    grade: str,
    load_kg: float,
    length_m: float,
    temp_source: str = "Встроенные данные (без ОГЗ)",
    b_mm: float = None,
    h_mm: float = None,
    t_mm: float = None,
    s_mm: float = None,
    m_kgm: float = None,
    geom_df: pd.DataFrame = None,
    comparison: list = None,
) -> bytes:
    """Возвращает байты .xlsx с данными и нативными Excel-графиками."""
    wb = openpyxl.Workbook()

    # ── Лист 1: Сводка ────────────────────────────────────────────────────────
    ws_sum = wb.active
    ws_sum.title = "Сводка"

    limit = res.fire_limit_minute
    cap0  = res.load_capacity["Несущая способность, кНм"].iloc[0]
    mom   = res.applied_moment_value
    if limit is None:
        limit_str = f"> {int(res.load_capacity['Время, мин'].iloc[-1])} мин"
    elif limit == 0:
        limit_str = "< 1 мин"
    else:
        limit_str = f"{limit} мин"

    summary_rows = [
        ("Нормативный документ",               doc_name),
        ("Профиль двутавра",                    profile_key),
        ("Марка стали",                         grade),
        ("Сосредоточенная нагрузка P, кг",      load_kg),
        ("Длина пролёта L, м",                  length_m),
        ("Источник температур",                 temp_source),
        ("", ""),
        ("РЕЗУЛЬТАТЫ", ""),
        ("Предел огнестойкости",                limit_str),
        ("Несущая способность при t = 0, кНм",  round(cap0, 2)),
        ("Момент от нагрузки M, кНм",           round(mom, 3)),
    ]
    if b_mm is not None:
        summary_rows = summary_rows[:6] + [
            ("Высота сечения h, мм",  h_mm),
            ("Ширина полки b, мм",    b_mm),
            ("Толщина полки t, мм",   t_mm),
            ("Толщина стенки s, мм",  s_mm),
        ] + summary_rows[6:]
    if m_kgm is not None:
        summary_rows.insert(6, ("Погонная масса, кг/м", m_kgm))

    for i, (lbl, val) in enumerate(summary_rows, 1):
        c1 = ws_sum.cell(row=i, column=1, value=lbl)
        c2 = ws_sum.cell(row=i, column=2, value=val)
        c1.font = _BODY_FONT
        c2.font = _BODY_FONT
        if lbl in ("РЕЗУЛЬТАТЫ",):
            c1.font = Font(name="Calibri", bold=True, size=11, color="1F3864")
        if lbl.startswith("Предел"):
            for c in (c1, c2):
                c.fill  = PatternFill("solid", fgColor="D6E4F0")
                c.font  = Font(name="Calibri", bold=True, size=10)
        c1.border = _BORDER
        c2.border = _BORDER
        c1.alignment = _CENTER
        c2.alignment = _CENTER

    ws_sum.column_dimensions["A"].width = 38
    ws_sum.column_dimensions["B"].width = 22

    # ── Лист 2: Несущая способность + График ─────────────────────────────────
    ws_cap = wb.create_sheet("Несущая способность")

    cap_df = pd.concat([
        res.load_capacity.set_index("Время, мин"),
        res.applied_moment.set_index("Время, мин"),
    ], axis=1).reset_index()
    cap_df.columns = ["Время, мин", "Несущая способность, кНм", "Момент от нагрузки, кНм"]

    n_cap = len(cap_df)
    _xl_write_df(ws_cap, cap_df)

    chart = _make_capacity_chart(ws_cap, n_cap)
    ws_cap.add_chart(chart, f"E2")

    # ── Лист 3: Температуры и прочность ──────────────────────────────────────
    ws_t = wb.create_sheet("Температуры и прочность")
    _xl_write_df(ws_t, res.strength_table.round(4))

    # ── Лист 4: Нейтральная ось и Rn ─────────────────────────────────────────
    ws_no = wb.create_sheet("НО и Rn")
    next_row = _xl_write_df(ws_no, res.compressed_zone.round(4))
    _xl_section_title(ws_no, "Нормативное сопротивление стали, кгс/см²",
                      next_row + 1, len(res.normative_resistance.columns))
    _xl_write_df(ws_no, res.normative_resistance.round(4), start_row=next_row + 2)

    # ── Лист 5: Усилия и плечи ────────────────────────────────────────────────
    ws_ef = wb.create_sheet("Усилия и плечи")
    next_row = _xl_write_df(ws_ef, res.efforts.round(4))
    _xl_section_title(ws_ef, "Плечи равнодействующих сил, мм",
                      next_row + 1, len(res.lever_arms.columns))
    _xl_write_df(ws_ef, res.lever_arms.round(4), start_row=next_row + 2)

    # ── Лист 6: Изгибающие моменты ───────────────────────────────────────────
    ws_m = wb.create_sheet("Изгибающие моменты")
    _xl_write_df(ws_m, res.bending_moments.round(4))

    # ── Лист 7: Параметры огнезащиты (если есть) ─────────────────────────────
    if geom_df is not None:
        ws_g = wb.create_sheet("Параметры огнезащиты")
        _xl_write_df(ws_g, geom_df)

    # ── Лист 8: Сравнение вариантов ОГЗ (если передано) ──────────────────────
    if comparison:
        ws_cmp = wb.create_sheet("Сравнение ОГЗ")

        t_max = max(int(r.load_capacity["Время, мин"].iloc[-1]) for _, r, _ in comparison)
        time_full = np.arange(0, t_max + 1, dtype=float)
        cmp_df = pd.DataFrame({"Время, мин": time_full})
        colors = []
        for label, res_cmp, color in comparison:
            cap_series = res_cmp.load_capacity.set_index("Время, мин")["Несущая способность, кНм"]
            cmp_df[label] = cap_series.reindex(time_full).values
            colors.append(color.lstrip("#").upper())
        cmp_df["Момент от нагрузки, кНм"] = comparison[0][1].applied_moment_value

        n_cmp = len(cmp_df)
        _xl_write_df(ws_cmp, cmp_df.round(4))

        chart_cmp = _make_comparison_chart_xl(ws_cmp, n_cmp, colors)
        ws_cmp.add_chart(chart_cmp, "G2")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
